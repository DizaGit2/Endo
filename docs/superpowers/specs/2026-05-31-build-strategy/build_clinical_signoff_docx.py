# -*- coding: utf-8 -*-
"""
Genera `lumen-revision-clinica-especialista.docx` — la versión en español, para un
especialista médico sin perfil técnico, del paquete de sign-off clínico
(`clinical-signoff-pack.md`, C-01 … C-15).

Uso (desde la raíz del repo):
    python docs/superpowers/specs/2026-05-31-build-strategy/build_clinical_signoff_docx.py

Requiere `python-docx` (pip install python-docx). El contenido clínico (valores,
fuentes) se copia tal cual del pack en inglés; aquí solo se reordena y se explica
en lenguaje llano. Si cambia un valor PO-interim en el pack, actualícelo también aquí.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from clinical_signoff_content import (
    CLOSING,
    INTRO,
    ITEMS,
    META,
    SECTIONS,
)

# ---------------------------------------------------------------------------
# Paleta (tokens del design system de Lumen, tema claro)
# ---------------------------------------------------------------------------
INK = RGBColor(0x3B, 0x2A, 0x20)
MUTED = RGBColor(0x8A, 0x6F, 0x5E)
ACCENT = RGBColor(0xC2, 0x5A, 0x36)
SAGE = RGBColor(0x7B, 0x8F, 0x6B)
ACCENT_SOFT_HEX = "F3D9CC"
SAGE_SOFT_HEX = "E4EADD"
INPUT_HEX = "FAF6EF"
BORDER_HEX = "D9D2C8"
PHASE_LUTEAL_HEX = "EEEDFE"

OUT_NAME = "lumen-revision-clinica-especialista.docx"


# ---------------------------------------------------------------------------
# Helpers de bajo nivel (XML de Word)
# ---------------------------------------------------------------------------
def shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def set_table_borders(table, color_hex: str = BORDER_HEX, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tbl_pr.append(borders)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def keep_with_next(paragraph, on: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = on


def set_cant_split(row) -> None:
    """La fila no se parte entre dos páginas."""
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def keep_row_with_next(row) -> None:
    """Truco de Word: 'conservar con el siguiente' en todos los párrafos de la fila."""
    for cell in row.cells:
        for p in cell.paragraphs:
            keep_with_next(p)


_NUM_CACHE: dict = {}


def restart_numbered(doc, paragraph) -> None:
    """Asigna al párrafo una numeración nueva (reinicia en 1) basada en 'List Number'."""
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    if "abstract" not in _NUM_CACHE:
        style = doc.styles["List Number"]
        num_pr = style.element.pPr.numPr if style.element.pPr is not None else None
        if num_pr is None or num_pr.numId is None:
            _NUM_CACHE["abstract"] = None
        else:
            num = numbering.num_having_numId(num_pr.numId.val)
            _NUM_CACHE["abstract"] = num.abstractNumId.val
    abstract_id = _NUM_CACHE["abstract"]
    if abstract_id is None:
        return
    if _NUM_CACHE.get("current") is None:
        new_num = numbering.add_num(abstract_id)
        new_num.add_lvlOverride(ilvl=0).add_startOverride(1)
        _NUM_CACHE["current"] = new_num.numId
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = _NUM_CACHE["current"]


def end_numbered_list() -> None:
    _NUM_CACHE["current"] = None


def numbered(doc, items: list[str], size: int | None = None):
    end_numbered_list()
    for text in items:
        p = doc.add_paragraph(style="List Number")
        restart_numbered(doc, p)
        add_rich(p, text, size=size)
        p.paragraph_format.space_after = Pt(3)
    end_numbered_list()


def set_lang(run_or_style, lang: str = "es-MX") -> None:
    rpr = run_or_style.element.get_or_add_rPr()
    el = rpr.find(qn("w:lang"))
    if el is None:
        el = OxmlElement("w:lang")
        rpr.append(el)
    el.set(qn("w:val"), lang)


def add_field(paragraph, instr: str) -> None:
    """Inserta un campo de Word (PAGE, NUMPAGES, …)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr_el, fld_sep, txt, fld_end):
        run._r.append(el)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


# ---------------------------------------------------------------------------
# Helpers de contenido
# ---------------------------------------------------------------------------
def add_rich(paragraph, text: str, size: int | None = None, color: RGBColor | None = None):
    """Texto con **negrita** inline (marcador simple de dos asteriscos)."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = bool(i % 2)
        if size:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
    return paragraph


def para(doc, text: str = "", style: str | None = None, size: int | None = None,
         color: RGBColor | None = None, space_after: int = 6, italic: bool = False,
         align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if text:
        add_rich(p, text, size=size, color=color)
        if italic:
            for r in p.runs:
                r.italic = True
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def bullet(doc, text: str, level: int = 0, size: int | None = None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    add_rich(p, text, size=size)
    p.paragraph_format.space_after = Pt(3)
    return p


def label(doc, text: str, color: RGBColor = ACCENT):
    """Etiqueta de subsección dentro de un ítem (p. ej. '¿Para qué se usa en la app?')."""
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    keep_with_next(p)
    return p


def simple_table(doc, header: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None,
                 header_fill: str = ACCENT_SOFT_HEX, font_size: int = 9, zebra: bool = True):
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, h in enumerate(header):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        shade_cell(cell, header_fill)
        set_cell_margins(cell)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich(p, val, size=font_size)
            set_cell_margins(cell)
            if zebra and ri % 2 == 1:
                shade_cell(cell, INPUT_HEX)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    n_rows = len(table.rows)
    for ri, row in enumerate(table.rows):
        set_cant_split(row)
        # Tablas cortas: enteras en una página. Largas: cabecera + 2 filas juntas.
        if (n_rows <= 9 and ri < n_rows - 1) or ri < 3:
            keep_row_with_next(row)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def callout(doc, title: str, body_lines: list[str], fill: str = SAGE_SOFT_HEX,
            title_color: RGBColor = SAGE, width_cm: float = 16.5):
    """Recuadro sombreado de una celda (para 'Pregunta para usted' y avisos)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color_hex=fill)
    cell = table.rows[0].cells[0]
    set_cant_split(table.rows[0])
    cell.width = Cm(width_cm)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = title_color
    p.paragraph_format.space_after = Pt(3)
    for line in body_lines:
        q = cell.add_paragraph()
        add_rich(q, line, size=10.5)
        q.paragraph_format.space_after = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def question_block(doc, question_lines: list[str], width_cm: float = 16.5):
    """Pregunta para el especialista + bloque de respuesta (casillas, comentario, firma),
    en una sola tabla cuyas filas no se parten ni se separan entre páginas."""
    table = doc.add_table(rows=4, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color_hex=BORDER_HEX)
    for row in table.rows:
        row.cells[0].width = Cm(width_cm)
        set_cell_margins(row.cells[0], top=100, bottom=100, left=180, right=180)
        set_cant_split(row)

    q = table.rows[0].cells[0]
    shade_cell(q, SAGE_SOFT_HEX)
    q.text = ""
    p = q.paragraphs[0]
    r = p.add_run("PREGUNTA PARA USTED")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = SAGE
    p.paragraph_format.space_after = Pt(3)
    for line in question_lines:
        qq = q.add_paragraph()
        add_rich(qq, line, size=10.5)
        qq.paragraph_format.space_after = Pt(3)

    c0 = table.rows[1].cells[0]
    c0.text = ""
    p = c0.paragraphs[0]
    r = p.add_run("SU RESPUESTA   ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = ACCENT
    for opt in ("De acuerdo", "Modificar", "No estoy de acuerdo"):
        box = p.add_run("☐ ")
        box.font.name = "Segoe UI Symbol"
        box._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI Symbol")
        box.font.size = Pt(12)
        t = p.add_run(opt + "     ")
        t.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(0)
    shade_cell(c0, INPUT_HEX)

    c1 = table.rows[2].cells[0]
    c1.text = ""
    p = c1.paragraphs[0]
    add_rich(p, "**Comentario, valor o redacción que propone (y la fuente, si la tiene):**", size=9.5)
    p.paragraph_format.space_after = Pt(2)
    for _ in range(3):
        q = c1.add_paragraph()
        q.paragraph_format.space_after = Pt(0)
        q.paragraph_format.space_before = Pt(8)
        rr = q.add_run("_" * 92)
        rr.font.color.rgb = RGBColor(0xB0, 0xA4, 0x98)
        rr.font.size = Pt(9)

    c2 = table.rows[3].cells[0]
    c2.text = ""
    p = c2.paragraphs[0]
    add_rich(p, "**Nombre y firma:** ______________________________________      **Fecha:** ____ / ____ / ________", size=9.5)
    p.paragraph_format.space_after = Pt(0)
    shade_cell(c2, INPUT_HEX)

    for row in table.rows[:-1]:
        keep_row_with_next(row)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ---------------------------------------------------------------------------
# Documento
# ---------------------------------------------------------------------------
def setup_document() -> Document:
    doc = Document()

    # Página carta (es-MX), márgenes cómodos para escribir a mano.
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    set_lang(normal)

    for name, size, color, before, after in (
        ("Title", 24, INK, 0, 6),
        ("Heading 1", 16, ACCENT, 18, 6),
        ("Heading 2", 13, INK, 16, 4),
        ("Heading 3", 11, INK, 10, 3),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        set_lang(st)
    for name in ("List Bullet", "List Bullet 2"):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        set_lang(st)

    # Encabezado y pie
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hr = hp.add_run("Lumen · Revisión clínica para el especialista · Confidencial")
    hr.font.size = Pt(8.5)
    hr.font.color.rgb = MUTED
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Página ")
    fr.font.size = Pt(9)
    fr.font.color.rgb = MUTED
    add_field(fp, "PAGE")
    fr2 = fp.add_run(" de ")
    fr2.font.size = Pt(9)
    fr2.font.color.rgb = MUTED
    add_field(fp, "NUMPAGES")
    return doc


def build_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Lumen")
    r.font.size = Pt(13)
    r.font.color.rgb = ACCENT
    r.bold = True
    p.paragraph_format.space_after = Pt(0)

    t = doc.add_paragraph(style="Title")
    t.add_run(META["title"])
    t.paragraph_format.space_after = Pt(2)

    st = para(doc, META["subtitle"], size=13, color=MUTED, space_after=14)

    meta_rows = [
        ["Dirigido a", META["for"]],
        ["Preparado por", META["from"]],
        ["Fecha", META["date"]],
        ["Versión", META["version"]],
        ["Estado", META["status"]],
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    set_table_borders(table, color_hex=BORDER_HEX)
    for k, v in meta_rows:
        cells = table.add_row().cells
        cells[0].width = Cm(3.6)
        cells[1].width = Cm(12.9)
        cells[0].text = ""
        cells[1].text = ""
        rk = cells[0].paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        rk.font.color.rgb = MUTED
        add_rich(cells[1].paragraphs[0], v, size=10)
        shade_cell(cells[0], INPUT_HEX)
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    for block in INTRO:
        kind = block[0]
        if kind == "h":
            doc.add_heading(block[1], level=1)
        elif kind == "p":
            para(doc, block[1])
        elif kind == "bullets":
            for b in block[1]:
                bullet(doc, b)
        elif kind == "numbered":
            numbered(doc, block[1])
        elif kind == "callout":
            callout(doc, block[1], block[2], fill=ACCENT_SOFT_HEX, title_color=ACCENT)
        elif kind == "table":
            simple_table(doc, block[1], block[2], col_widths_cm=block[3] if len(block) > 3 else None)


def build_item(doc: Document, item: dict) -> None:
    h = doc.add_heading(level=2)
    r0 = h.add_run(item["code"] + "  ")
    r0.font.color.rgb = ACCENT
    h.add_run(item["title"])

    if item.get("tags"):
        p = doc.add_paragraph()
        r = p.add_run("  ·  ".join(item["tags"]))
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        r.italic = True
        p.paragraph_format.space_after = Pt(2)

    label(doc, "¿Para qué se usa en la app?")
    for line in item["usage"]:
        para(doc, line, space_after=4)

    # C-16 has no PO decision — the product owner deliberately declined to invent
    # one — so an item may override this heading. Everything else is unchanged.
    label(doc, item.get("decision_label",
                        "Decisión provisional de Carolina (responsable del producto)"))
    for block in item["decision"]:
        if isinstance(block, str):
            para(doc, block, space_after=4)
        elif block[0] == "bullets":
            for b in block[1]:
                bullet(doc, b)
        elif block[0] == "numbered":
            numbered(doc, block[1])
        elif block[0] == "table":
            simple_table(doc, block[1], block[2], col_widths_cm=block[3] if len(block) > 3 else None,
                         font_size=block[4] if len(block) > 4 else 9)
        elif block[0] == "note":
            callout(doc, "Nota", [block[1]], fill=PHASE_LUTEAL_HEX, title_color=MUTED)
        elif block[0] == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            add_rich(p, block[1], size=10)
            for rr in p.runs:
                rr.italic = True
            p.paragraph_format.space_after = Pt(4)
    if item.get("sources"):
        p = doc.add_paragraph()
        r = p.add_run("En qué se basa: ")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        r2 = p.add_run(item["sources"])
        r2.font.size = Pt(9)
        r2.font.color.rgb = MUTED
        p.paragraph_format.space_after = Pt(4)

    question_block(doc, item["question"])


def build_closing(doc: Document) -> None:
    h = doc.add_heading(CLOSING["summary_title"], level=1)
    h.paragraph_format.page_break_before = True  # sin página en blanco si la anterior quedó llena
    para(doc, CLOSING["summary_intro"])
    header = ["Punto", "Tema", "De acuerdo", "Modificar", "No de acuerdo", "Observaciones"]
    rows = []
    for it in ITEMS:
        rows.append([it["code"], it["short"], "☐", "☐", "☐", ""])
    table = simple_table(doc, header, rows, col_widths_cm=[1.3, 5.4, 1.9, 1.9, 2.1, 4.4], zebra=True, font_size=9)
    # Casillas en fuente con el glifo
    for row in table.rows[1:]:
        for ci in (2, 3, 4):
            for p in row.cells[ci].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = "Segoe UI Symbol"
                    r.font.size = Pt(12)

    for block in CLOSING["blocks"]:
        kind = block[0]
        if kind == "h":
            doc.add_heading(block[1], level=1)
        elif kind == "p":
            para(doc, block[1])
        elif kind == "bullets":
            for b in block[1]:
                bullet(doc, b)

    # Firma final
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    set_table_borders(table, color_hex=BORDER_HEX)
    labels = [["Nombre completo del especialista", "Especialidad / institución"],
              ["Firma", "Fecha"]]
    for ri in range(2):
        for ci in range(2):
            cell = table.rows[ri].cells[ci]
            cell.width = Cm(8.25)
            cell.text = ""
            rr = cell.paragraphs[0].add_run(labels[ri][ci])
            rr.font.size = Pt(9)
            rr.font.color.rgb = MUTED
            for _ in range(2):
                cell.add_paragraph()
            set_cell_margins(cell, top=100, bottom=100, left=160, right=160)


def main(out_dir: Path) -> Path:
    doc = setup_document()
    build_cover(doc)

    for si, section in enumerate(SECTIONS):
        h = doc.add_heading(section["title"], level=1)
        if si == 0:
            h.paragraph_format.page_break_before = True
        para(doc, section["intro"], color=MUTED, size=10.5, space_after=8)
        for item in ITEMS:
            if item["section"] == section["key"]:
                build_item(doc, item)

    build_closing(doc)

    out = out_dir / OUT_NAME
    doc.save(out)
    return out


if __name__ == "__main__":
    here = Path(__file__).resolve().parent
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else here
    result = main(target)
    print(f"Escrito: {result}")
